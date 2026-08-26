from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

MD = r"D:\E Drive\Engineering\testleaf\workspace\autoFrameX\docs\INTERVIEW_GUIDE.md"
OUT = r"D:\E Drive\Engineering\testleaf\workspace\autoFrameX\docs\autoFrameX_Interview_Guide.docx"

doc = Document()

# ── page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin   = Inches(1.1)
    section.right_margin  = Inches(1.1)

# ── style helpers ─────────────────────────────────────────────────────────────
def set_font(run, name="Calibri", size=11, bold=False, italic=False, color=None):
    run.font.name  = name
    run.font.size  = Pt(size)
    run.font.bold  = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_heading(text, level):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.space_before = Pt(10 if level == 1 else 6)
    p.paragraph_format.space_after  = Pt(4)
    for run in p.runs:
        run.font.name = "Calibri"
        if level == 1:
            run.font.size  = Pt(18)
            run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
        elif level == 2:
            run.font.size  = Pt(14)
            run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
        else:
            run.font.size  = Pt(12)
            run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    return p

def add_normal(text, bold=False, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_font(run, bold=bold, italic=italic)
    return p

def add_bullet(text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent  = Inches(0.25 * (level + 1))
    p.paragraph_format.space_after  = Pt(2)
    # inline code inside bullet
    parts = re.split(r'(`[^`]+`)', text)
    for part in parts:
        if part.startswith('`') and part.endswith('`'):
            run = p.add_run(part[1:-1])
            set_font(run, name="Courier New", size=9.5)
        else:
            # handle **bold**
            bold_parts = re.split(r'(\*\*[^*]+\*\*)', part)
            for bp in bold_parts:
                if bp.startswith('**') and bp.endswith('**'):
                    run = p.add_run(bp[2:-2])
                    set_font(run, bold=True)
                else:
                    run = p.add_run(bp)
                    set_font(run)
    return p

def add_code_block(lines):
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_after  = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.left_indent  = Inches(0.3)
        # shade background
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'),   'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'),  'F2F2F2')
        pPr.append(shd)
        run = p.add_run(line if line else " ")
        set_font(run, name="Courier New", size=9)

def add_table(headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    # header row
    hrow = t.rows[0]
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            set_font(run, bold=True, size=10)
        # header bg
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'),   'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'),  'BDD7EE')
        tcPr.append(shd)
    # data rows
    for ri, row in enumerate(rows):
        drow = t.rows[ri + 1]
        for ci, cell_text in enumerate(row):
            cell = drow.cells[ci]
            cell.text = ""
            p = cell.paragraphs[0]
            parts = re.split(r'(`[^`]+`)', cell_text)
            for part in parts:
                if part.startswith('`') and part.endswith('`'):
                    run = p.add_run(part[1:-1])
                    set_font(run, name="Courier New", size=9)
                else:
                    run = p.add_run(part)
                    set_font(run, size=10)
            if ri % 2 == 1:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'),   'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'),  'F5F9FF')
                tcPr.append(shd)
    doc.add_paragraph()

def add_inline_para(text):
    """Paragraph with inline `code`, **bold**, and *italic* support."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    # split on backtick code, bold, italic
    token_re = re.compile(r'(`[^`]+`|\*\*[^*]+\*\*|\*[^*]+\*)')
    parts = token_re.split(text)
    for part in parts:
        if part.startswith('`') and part.endswith('`'):
            run = p.add_run(part[1:-1])
            set_font(run, name="Courier New", size=9.5)
        elif part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            set_font(run, bold=True)
        elif part.startswith('*') and part.endswith('*'):
            run = p.add_run(part[1:-1])
            set_font(run, italic=True)
        else:
            run = p.add_run(part)
            set_font(run)
    return p

def add_qa(question, answer):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(2)
    r = p.add_run(question)
    set_font(r, bold=True, color=(0x1F, 0x49, 0x7D))
    add_inline_para(answer)

# ── parse and render ──────────────────────────────────────────────────────────
with open(MD, encoding="utf-8") as f:
    lines = f.readlines()

# title page
doc.add_paragraph()
tp = doc.add_paragraph()
tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = tp.add_run("autoFrameX")
set_font(r, size=28, bold=True, color=(0x1F, 0x49, 0x7D))

tp2 = doc.add_paragraph()
tp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = tp2.add_run("Interview Guide")
set_font(r2, size=18, color=(0x2E, 0x74, 0xB5))

tp3 = doc.add_paragraph()
tp3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = tp3.add_run("Enterprise Test Automation Framework — Java 17 · Selenium 4 · TestNG 7")
set_font(r3, size=11, italic=True, color=(0x60, 0x60, 0x60))

doc.add_page_break()

i = 0
in_code  = False
in_table = False
code_lines   = []
table_headers = []
table_rows    = []

def flush_table():
    global in_table, table_headers, table_rows
    if table_headers:
        add_table(table_headers, table_rows)
    in_table = False
    table_headers = []
    table_rows    = []

def flush_code():
    global in_code, code_lines
    add_code_block(code_lines)
    doc.add_paragraph()
    in_code    = False
    code_lines = []

while i < len(lines):
    raw  = lines[i].rstrip('\n')
    line = raw.strip()

    # ── code fence ────────────────────────────────────────────────────────────
    if line.startswith('```'):
        if in_table:
            flush_table()
        if not in_code:
            in_code    = True
            code_lines = []
        else:
            flush_code()
        i += 1
        continue

    if in_code:
        code_lines.append(raw)
        i += 1
        continue

    # ── table row ─────────────────────────────────────────────────────────────
    if line.startswith('|'):
        cells = [c.strip() for c in line.strip('|').split('|')]
        if all(re.match(r'^[-: ]+$', c) for c in cells):
            i += 1
            continue
        if not in_table:
            in_table      = True
            table_headers = cells
            table_rows    = []
        else:
            table_rows.append(cells)
        i += 1
        continue
    else:
        if in_table:
            flush_table()

    # ── headings ──────────────────────────────────────────────────────────────
    if line.startswith('# ') and not line.startswith('## '):
        add_heading(line[2:], 1)
        i += 1; continue
    if line.startswith('## '):
        add_heading(line[3:], 2)
        i += 1; continue
    if line.startswith('### '):
        add_heading(line[4:], 3)
        i += 1; continue

    # ── horizontal rule ───────────────────────────────────────────────────────
    if line in ('---', '***', '___'):
        doc.add_paragraph()
        i += 1; continue

    # ── blockquote ────────────────────────────────────────────────────────────
    if line.startswith('> '):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent  = Inches(0.4)
        p.paragraph_format.space_after  = Pt(6)
        run = p.add_run(line[2:])
        set_font(run, italic=True, color=(0x40, 0x40, 0x40))
        i += 1; continue

    # ── bullet ────────────────────────────────────────────────────────────────
    m = re.match(r'^(\s*)([-*+]|\d+\.) (.+)', line)
    if m:
        indent = len(m.group(1)) // 2
        add_bullet(m.group(3), level=indent)
        i += 1; continue

    # ── Q&A pattern ───────────────────────────────────────────────────────────
    if line.startswith('**Q:') and line.endswith('**'):
        question = line[2:-2]
        # next non-empty line is the answer
        j = i + 1
        while j < len(lines) and not lines[j].strip():
            j += 1
        answer = ""
        if j < len(lines):
            ans_line = lines[j].strip()
            if ans_line.startswith('A:'):
                answer = ans_line[2:].strip()
                i = j + 1
            else:
                i += 1
        else:
            i += 1
        add_qa(question, answer)
        continue

    # ── blank line ────────────────────────────────────────────────────────────
    if not line:
        i += 1; continue

    # ── normal paragraph ──────────────────────────────────────────────────────
    add_inline_para(line)
    i += 1

if in_table:
    flush_table()
if in_code:
    flush_code()

doc.save(OUT)
print(f"Saved: {OUT}")
