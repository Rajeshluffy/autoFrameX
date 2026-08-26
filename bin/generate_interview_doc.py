import sys
try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

doc = Document()

# ── Helper functions ──────────────────────────────────────────────────────────
def set_heading_color(para, rgb):
    for run in para.runs:
        run.font.color.rgb = RGBColor(*rgb)

def add_title(doc, text):
    p = doc.add_heading(text, level=0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.font.size = Pt(22)
        run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    return p

def add_section_heading(doc, text):
    p = doc.add_heading(text, level=1)
    for run in p.runs:
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    return p

def add_question_heading(doc, text):
    p = doc.add_heading(text, level=2)
    for run in p.runs:
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    return p

def add_para(doc, text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    return p

def add_star_para(doc, label, text):
    p = doc.add_paragraph()
    r1 = p.add_run(label + " ")
    r1.font.bold = True
    r1.font.name = "Calibri"
    r1.font.size = Pt(11)
    r1.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    r2 = p.add_run(text)
    r2.font.name = "Calibri"
    r2.font.size = Pt(11)
    return p

def add_code(doc, code_text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(code_text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), "F2F2F2")
    p._p.get_or_add_pPr().append(shading)
    return p

def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    return p

def page_break(doc):
    doc.add_page_break()

# ── TITLE PAGE ────────────────────────────────────────────────────────────────
add_title(doc, "Interview Questions & Answers")
add_title(doc, "STAR Technique - Senior QA/SDET")
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Projects: autoFrameX | MS Dynamics 365 | Salesforce | TransUnion | PayPal")
run.font.name = "Calibri"
run.font.size = Pt(12)
run.font.italic = True
run.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)
doc.add_paragraph()
p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run("Experience: 5+ Years | Stack: Java, Selenium, TestNG, RestAssured, Jenkins, JMeter")
r2.font.name = "Calibri"
r2.font.size = Pt(11)
r2.font.color.rgb = RGBColor(0x59, 0x59, 0x59)
page_break(doc)

# ── TABLE OF CONTENTS ─────────────────────────────────────────────────────────
add_section_heading(doc, "Table of Contents")
for item in [
    "Section 1: Behavioral / HR Questions",
    "Section 2: Java Core",
    "Section 3: Selenium",
    "Section 4: API Testing (RestAssured)",
    "Section 5: TestNG",
    "Section 6: CI/CD & Git",
    "Section 7: Design Patterns",
    "Section 8: SQL",
    "Section 9: DSA / Coding Problems",
]:
    add_bullet(doc, item)
page_break(doc)

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 1: BEHAVIORAL / HR
# ═══════════════════════════════════════════════════════════════════════════════
add_section_heading(doc, "Section 1: Behavioral / HR Questions")

add_question_heading(doc, "Q1. Tell me about yourself and your roles & responsibilities")
add_para(doc, "Answer:", bold=True)
add_star_para(doc, "Situation:", "I am a Senior QA/SDET with 5+ years of experience in test automation across enterprise web applications, APIs, and performance testing.")
add_star_para(doc, "Task:", "My core responsibility is to design, build, and maintain scalable automation frameworks that reduce manual effort and accelerate release cycles.")
add_star_para(doc, "Action:",
    "In my current role, I architected autoFrameX - a hybrid Selenium+TestNG+RestAssured framework with Page Object Model, "
    "data-driven testing, parallel execution via Selenium Grid, and CI/CD integration through Jenkins. "
    "I have worked on enterprise applications including MS Dynamics 365 (CRM), Salesforce (low-code CRM), "
    "TransUnion (credit bureau data), and PayPal (payment gateway). "
    "Day-to-day I write automation scripts, review test plans, triage failures in Jenkins pipelines, "
    "collaborate with developers on bug fixes, and mentor junior QA engineers.")
add_star_para(doc, "Result:",
    "Achieved 80%+ automation coverage for regression suites, reduced regression cycle from 3 days to 4 hours, "
    "and improved defect detection rate by 35% through early shift-left testing.")

add_question_heading(doc, "Q2. How do you help developers in your project?")
add_para(doc, "Answer:", bold=True)
add_star_para(doc, "Situation:", "Developers often push code without knowing the downstream impact on existing functionality.")
add_star_para(doc, "Task:", "My role is to act as a quality gate and a collaborative partner - not just a tester.")
add_star_para(doc, "Action:",
    "1. Share automation test results from Jenkins after every build so developers get instant feedback.\n"
    "2. Write detailed bug reports with steps to reproduce, screenshots, and logs.\n"
    "3. Participate in code reviews to flag testability issues early.\n"
    "4. Create reusable utility methods (waitForElement, scrollToElement) that developers can also use.\n"
    "5. During sprint planning, help developers understand acceptance criteria and edge cases.")
add_star_para(doc, "Result:", "Reduced back-and-forth on bug reports by 40% and improved developer confidence in releases.")

add_question_heading(doc, "Q3. How do you handle tight deadlines with many scenarios to complete?")
add_para(doc, "Answer:", bold=True)
add_star_para(doc, "Situation:", "During a major MS Dynamics 365 release, we had 200+ test scenarios to execute in 2 days before go-live.")
add_star_para(doc, "Task:", "Complete regression testing without compromising quality.")
add_star_para(doc, "Action:",
    "1. Prioritized test cases using risk-based testing - P1 (critical business flows) first.\n"
    "2. Ran automation suite in parallel using Selenium Grid (4 nodes) to cut execution time by 75%.\n"
    "3. Delegated manual exploratory testing of new features to junior QA.\n"
    "4. Communicated daily status to the PM with pass/fail counts and risk areas.\n"
    "5. Deferred low-risk edge cases to the next sprint with documented risk acceptance.")
add_star_para(doc, "Result:", "Completed 95% of critical scenarios on time. Go-live was successful with zero P1 defects in production.")

add_question_heading(doc, "Q4. How do you help team members?")
add_para(doc, "Answer:", bold=True)
add_star_para(doc, "Situation:", "A junior QA engineer on my team was struggling with XPath locators and understanding the POM structure.")
add_star_para(doc, "Task:", "Upskill the team member without impacting sprint velocity.")
add_star_para(doc, "Action:",
    "1. Conducted 1-on-1 knowledge transfer sessions on XPath strategies and POM design.\n"
    "2. Created a wiki with framework onboarding guide, coding standards, and common patterns.\n"
    "3. Did pair programming - reviewed their PRs and gave constructive feedback.\n"
    "4. Shared curated resources: Udemy courses, official Selenium docs, and internal framework README.")
add_star_para(doc, "Result:", "Within 3 weeks, the junior QA was independently writing automation scripts and contributing to the sprint.")

add_question_heading(doc, "Q5. How do you stay updated with the latest technology?")
add_para(doc, "Answer:", bold=True)
add_para(doc, "I follow a structured approach to continuous learning:")
add_bullet(doc, "Read official documentation and release notes (Selenium 4, TestNG, RestAssured)")
add_bullet(doc, "Follow thought leaders on LinkedIn and Medium (Angie Jones, Alan Richardson)")
add_bullet(doc, "Take Udemy/Pluralsight courses - recently completed Docker for QA and Playwright basics")
add_bullet(doc, "Contribute to and explore open-source projects on GitHub")
add_bullet(doc, "Attend webinars and QA conferences (SauceCon, TestBash)")
add_bullet(doc, "Practice DSA on LeetCode to sharpen coding skills")
add_bullet(doc, "Experiment in personal projects - e.g., added Cucumber BDD layer to autoFrameX")

add_question_heading(doc, "Q6. How do you understand a big system in 2-3 weeks?")
add_para(doc, "Answer:", bold=True)
add_star_para(doc, "Situation:", "When I joined the TransUnion project, the application had 50+ modules with complex data flows.")
add_star_para(doc, "Task:", "Ramp up quickly to start contributing to automation within 2 weeks.")
add_star_para(doc, "Action:",
    "1. Day 1-3: Read all available documentation - BRD, FRD, test plans, existing test cases.\n"
    "2. Day 4-7: Shadowed manual testers and attended sprint ceremonies to understand business flows.\n"
    "3. Day 8-10: Explored the application end-to-end, mapped user journeys, identified critical paths.\n"
    "4. Day 11-14: Set up the automation framework, wrote smoke tests for the top 5 critical flows.\n"
    "5. Asked targeted questions to SMEs - specific ones based on my research, not generic questions.")
add_star_para(doc, "Result:", "Was contributing automation scripts by end of week 2 and presented a test coverage plan by week 3.")

add_question_heading(doc, "Q7. What are your short-term and long-term goals?")
add_para(doc, "Answer:", bold=True)
add_para(doc, "Short-term (6-12 months):", bold=True)
add_bullet(doc, "Master Playwright and add it as an alternative UI automation layer in autoFrameX")
add_bullet(doc, "Get AWS Certified Developer Associate certification")
add_bullet(doc, "Deepen expertise in contract testing using Pact")
add_para(doc, "Long-term (2-5 years):", bold=True)
add_bullet(doc, "Grow into a QA Architect or Engineering Manager role")
add_bullet(doc, "Build a fully AI-assisted test generation pipeline integrated into CI/CD")
add_bullet(doc, "Contribute to open-source test automation tools")

add_question_heading(doc, "Q8. How do you handle conflicts with developers?")
add_para(doc, "Answer:", bold=True)
add_star_para(doc, "Situation:", "A developer disagreed with a P1 bug I raised, claiming it was working as designed - but it was breaking a critical payment flow in the PayPal integration.")
add_star_para(doc, "Task:", "Resolve the conflict professionally while ensuring the defect was addressed.")
add_star_para(doc, "Action:",
    "1. Scheduled a 15-minute call - not a long email chain.\n"
    "2. Demonstrated the bug with a recorded video and exact steps to reproduce.\n"
    "3. Referenced the acceptance criteria from the user story to show expected behavior.\n"
    "4. Involved the BA and PO to clarify the business requirement.\n"
    "5. Kept the conversation focused on the issue, not the person.")
add_star_para(doc, "Result:", "The developer acknowledged the bug, fixed it within the sprint, and we established a shared definition of done going forward.")

add_question_heading(doc, "Q9. What is your strength and weakness?")
add_para(doc, "Answer:", bold=True)
add_para(doc, "Strength:", bold=True)
add_para(doc,
    "My biggest strength is framework design and problem-solving. I can look at a complex testing challenge "
    "and architect a clean, maintainable solution. For example, I designed autoFrameX with a modular architecture "
    "- separate layers for UI, API, data, and reporting - which made it easy to onboard new team members and extend functionality.")
add_para(doc, "Weakness:", bold=True)
add_para(doc,
    "I used to over-engineer solutions - spending too much time making things perfect rather than good enough for now. "
    "I have been actively working on this by applying the YAGNI principle (You Aren't Gonna Need It) and timeboxing design decisions. "
    "This has improved my delivery speed without sacrificing quality.")

page_break(doc)

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 2: JAVA CORE
# ═══════════════════════════════════════════════════════════════════════════════
add_section_heading(doc, "Section 2: Java Core")

add_question_heading(doc, "Q1. Explain OOP concepts with framework examples")
add_para(doc, "Answer:", bold=True)
add_para(doc, "The four pillars of OOP and how they appear in autoFrameX:")
add_bullet(doc, "Encapsulation: Page classes encapsulate locators and actions. External code only calls methods like loginPage.login(user, pass) - it never touches WebElements directly.")
add_bullet(doc, "Inheritance: BasePage holds common methods (waitForElement, click, type). All page classes extend BasePage, inheriting driver and utility methods.")
add_bullet(doc, "Polymorphism: BrowserFactory.getDriver(browserName) returns a WebDriver reference. At runtime it could be ChromeDriver, FirefoxDriver, or EdgeDriver - same interface, different behavior.")
add_bullet(doc, "Abstraction: IPage interface defines the contract (initElements, isLoaded). Concrete page classes implement it. Callers depend on the interface, not the implementation.")
add_code(doc,
    "// Encapsulation\n"
    "public class LoginPage extends BasePage {\n"
    "    private By usernameField = By.id(\"username\");\n"
    "    public void login(String user, String pass) {\n"
    "        type(usernameField, user);\n"
    "    }\n"
    "}\n\n"
    "// Polymorphism\n"
    "WebDriver driver = BrowserFactory.getDriver(\"chrome\"); // returns ChromeDriver\n"
    "WebDriver driver = BrowserFactory.getDriver(\"firefox\"); // returns FirefoxDriver")

add_question_heading(doc, "Q2. Difference between Interface and Abstract class - where used in project")
add_para(doc, "Answer:", bold=True)
add_para(doc, "Interface:", bold=True)
add_bullet(doc, "All methods are abstract by default (Java 7); can have default/static methods (Java 8+)")
add_bullet(doc, "Supports multiple inheritance")
add_bullet(doc, "No instance variables (only public static final constants)")
add_bullet(doc, "Use when: defining a contract that unrelated classes must follow")
add_para(doc, "Abstract Class:", bold=True)
add_bullet(doc, "Can have both abstract and concrete methods")
add_bullet(doc, "Single inheritance only")
add_bullet(doc, "Can have instance variables and constructors")
add_bullet(doc, "Use when: sharing common code among closely related classes")
add_para(doc, "In autoFrameX:", bold=True)
add_bullet(doc, "Interface IPage: defines isLoaded() and getTitle() - all page objects implement this contract")
add_bullet(doc, "Abstract class BasePage: holds WebDriver instance, common wait methods, and click/type utilities - all page classes extend this")
add_code(doc,
    "public interface IPage {\n"
    "    boolean isLoaded();\n"
    "    String getTitle();\n"
    "}\n\n"
    "public abstract class BasePage implements IPage {\n"
    "    protected WebDriver driver;\n"
    "    protected WebDriverWait wait;\n"
    "    public BasePage(WebDriver driver) {\n"
    "        this.driver = driver;\n"
    "        this.wait = new WebDriverWait(driver, Duration.ofSeconds(10));\n"
    "    }\n"
    "    protected void click(By locator) {\n"
    "        wait.until(ExpectedConditions.elementToBeClickable(locator)).click();\n"
    "    }\n"
    "}")

add_question_heading(doc, "Q3. Polymorphism with framework example")
add_para(doc, "Answer:", bold=True)
add_para(doc,
    "Polymorphism means one interface, many implementations. In Java it comes in two forms:")
add_bullet(doc, "Compile-time (Method Overloading): same method name, different parameters")
add_bullet(doc, "Runtime (Method Overriding): child class overrides parent class method")
add_para(doc, "In autoFrameX - Runtime Polymorphism:", bold=True)
add_code(doc,
    "// BrowserFactory returns WebDriver (parent type)\n"
    "// but actual object is ChromeDriver/FirefoxDriver (child type)\n"
    "public class BrowserFactory {\n"
    "    public static WebDriver getDriver(String browser) {\n"
    "        switch(browser.toLowerCase()) {\n"
    "            case \"chrome\": return new ChromeDriver();\n"
    "            case \"firefox\": return new FirefoxDriver();\n"
    "            case \"edge\": return new EdgeDriver();\n"
    "            default: throw new IllegalArgumentException(\"Unknown browser: \" + browser);\n"
    "        }\n"
    "    }\n"
    "}\n"
    "// Usage - polymorphic reference\n"
    "WebDriver driver = BrowserFactory.getDriver(browserName);")

add_question_heading(doc, "Q4. Collections - List vs Set, HashMap internals")
add_para(doc, "Answer:", bold=True)
add_para(doc, "List vs Set:", bold=True)
add_bullet(doc, "List (ArrayList, LinkedList): ordered, allows duplicates, index-based access. Use when order matters or duplicates are valid.")
add_bullet(doc, "Set (HashSet, TreeSet, LinkedHashSet): no duplicates. HashSet is O(1) for add/contains. TreeSet is sorted (O(log n)). LinkedHashSet maintains insertion order.")
add_para(doc, "HashMap internals:", bold=True)
add_bullet(doc, "Backed by an array of buckets (default 16). Key's hashCode() determines bucket index.")
add_bullet(doc, "Collision handling: Java 8+ uses a linked list that converts to a Red-Black Tree when bucket size > 8 (O(n) -> O(log n)).")
add_bullet(doc, "Load factor 0.75: when 75% full, the map resizes (doubles) and rehashes all entries.")
add_bullet(doc, "Keys must implement hashCode() and equals() correctly.")
add_para(doc, "In autoFrameX:", bold=True)
add_bullet(doc, "Map<String, String> testData = new HashMap<>() - stores test data key-value pairs from Excel/JSON")
add_bullet(doc, "Set<String> windowHandles = driver.getWindowHandles() - returns Set of window handle strings")
add_code(doc,
    "// Collections in framework\n"
    "Map<String, String> config = new HashMap<>();\n"
    "config.put(\"baseUrl\", \"https://app.example.com\");\n"
    "config.put(\"browser\", \"chrome\");\n\n"
    "List<WebElement> rows = driver.findElements(By.tagName(\"tr\"));\n"
    "Set<String> handles = driver.getWindowHandles();")

add_question_heading(doc, "Q5. Exception handling - checked vs unchecked")
add_para(doc, "Answer:", bold=True)
add_bullet(doc, "Checked exceptions: Checked at compile time. Must be caught or declared with throws. Examples: IOException, SQLException. Extend Exception.")
add_bullet(doc, "Unchecked exceptions: Not checked at compile time. Extend RuntimeException. Examples: NullPointerException, ArrayIndexOutOfBoundsException, IllegalArgumentException.")
add_para(doc, "In autoFrameX:", bold=True)
add_code(doc,
    "// Checked - must handle\n"
    "try {\n"
    "    FileInputStream fis = new FileInputStream(\"config.properties\");\n"
    "} catch (FileNotFoundException e) {\n"
    "    log.error(\"Config file not found\", e);\n"
    "    throw new RuntimeException(\"Config file missing\", e);\n"
    "}\n\n"
    "// Custom unchecked exception in framework\n"
    "public class FrameworkException extends RuntimeException {\n"
    "    public FrameworkException(String message, Throwable cause) {\n"
    "        super(message, cause);\n"
    "    }\n"
    "}")

add_question_heading(doc, "Q6. String immutability")
add_para(doc, "Answer:", bold=True)
add_para(doc,
    "String is immutable in Java - once created, its value cannot be changed. "
    "Any operation that appears to modify a String actually creates a new String object in the String Pool.")
add_para(doc, "Why immutable?", bold=True)
add_bullet(doc, "Security: String is used for class loading, network connections, file paths - immutability prevents tampering")
add_bullet(doc, "Thread safety: immutable objects are inherently thread-safe")
add_bullet(doc, "String Pool: JVM can cache and reuse String literals, saving memory")
add_bullet(doc, "HashCode caching: String caches its hashCode, making it efficient as HashMap key")
add_para(doc, "For mutable strings use StringBuilder (single-threaded) or StringBuffer (thread-safe).", bold=False)
add_code(doc,
    "String s = \"hello\";\n"
    "s.concat(\" world\"); // creates new String, s still points to \"hello\"\n"
    "s = s.concat(\" world\"); // now s points to new String \"hello world\"\n\n"
    "// Mutable alternative\n"
    "StringBuilder sb = new StringBuilder(\"hello\");\n"
    "sb.append(\" world\"); // modifies same object - efficient in loops")

add_question_heading(doc, "Q7. final, finally, finalize")
add_para(doc, "Answer:", bold=True)
add_bullet(doc, "final: keyword. On variable: value cannot be reassigned. On method: cannot be overridden. On class: cannot be subclassed. Example: public static final String BASE_URL = \"https://api.example.com\";")
add_bullet(doc, "finally: block in try-catch-finally. Always executes regardless of exception. Used for cleanup (closing streams, drivers). Exception: System.exit() or JVM crash prevents it.")
add_bullet(doc, "finalize(): method in Object class. Called by GC before object is garbage collected. Deprecated in Java 9, removed in Java 18. Unreliable - use try-with-resources instead.")
add_code(doc,
    "// final in framework\n"
    "public static final int TIMEOUT = 10;\n\n"
    "// finally for driver cleanup\n"
    "try {\n"
    "    driver.get(url);\n"
    "    // test steps\n"
    "} catch (Exception e) {\n"
    "    log.error(\"Test failed\", e);\n"
    "} finally {\n"
    "    driver.quit(); // always runs\n"
    "}")

add_question_heading(doc, "Q8. Static keyword")
add_para(doc, "Answer:", bold=True)
add_bullet(doc, "Static variable: belongs to the class, not instances. Shared across all objects. Example: static int instanceCount.")
add_bullet(doc, "Static method: can be called without creating an object. Cannot access instance variables. Example: BrowserFactory.getDriver().")
add_bullet(doc, "Static class: only inner classes can be static. Static inner class does not hold reference to outer class.")
add_bullet(doc, "Static block: runs once when class is loaded. Used for one-time initialization.")
add_para(doc, "In autoFrameX - ThreadLocal + static for parallel execution:", bold=True)
add_code(doc,
    "public class DriverManager {\n"
    "    private static ThreadLocal<WebDriver> driverThread = new ThreadLocal<>();\n\n"
    "    public static WebDriver getDriver() {\n"
    "        return driverThread.get();\n"
    "    }\n"
    "    public static void setDriver(WebDriver driver) {\n"
    "        driverThread.set(driver);\n"
    "    }\n"
    "    public static void removeDriver() {\n"
    "        driverThread.remove();\n"
    "    }\n"
    "}")

add_question_heading(doc, "Q9. Multithreading and ThreadLocal - used in framework for parallel execution")
add_para(doc, "Answer:", bold=True)
add_para(doc,
    "In parallel test execution with TestNG, multiple threads run simultaneously. "
    "If all threads share a single WebDriver instance, they will interfere with each other. "
    "ThreadLocal solves this by giving each thread its own isolated copy of the WebDriver.")
add_bullet(doc, "ThreadLocal<T>: provides thread-local variables. Each thread has its own independent copy.")
add_bullet(doc, "set(value): stores value for current thread")
add_bullet(doc, "get(): retrieves value for current thread")
add_bullet(doc, "remove(): cleans up to prevent memory leaks")
add_code(doc,
    "// DriverManager.java - thread-safe driver management\n"
    "public class DriverManager {\n"
    "    private static final ThreadLocal<WebDriver> DRIVER = new ThreadLocal<>();\n\n"
    "    public static void initDriver(String browser) {\n"
    "        WebDriver driver = BrowserFactory.getDriver(browser);\n"
    "        DRIVER.set(driver);\n"
    "    }\n"
    "    public static WebDriver getDriver() { return DRIVER.get(); }\n"
    "    public static void quitDriver() {\n"
    "        if (DRIVER.get() != null) {\n"
    "            DRIVER.get().quit();\n"
    "            DRIVER.remove(); // prevent memory leak\n"
    "        }\n"
    "    }\n"
    "}\n\n"
    "// testng.xml - parallel execution config\n"
    "// <suite name=\"Suite\" parallel=\"tests\" thread-count=\"4\">")

add_question_heading(doc, "Q10. Java 8 features - streams, lambda")
add_para(doc, "Answer:", bold=True)
add_para(doc, "Key Java 8 features used in autoFrameX:", bold=True)
add_bullet(doc, "Lambda expressions: concise anonymous function syntax. (params) -> expression")
add_bullet(doc, "Stream API: functional-style operations on collections (filter, map, collect, reduce)")
add_bullet(doc, "Optional: container to avoid NullPointerException")
add_bullet(doc, "Default methods in interfaces: backward-compatible interface evolution")
add_bullet(doc, "Method references: shorthand for lambdas calling existing methods")
add_code(doc,
    "// Lambda + Stream in framework\n"
    "List<String> failedTests = testResults.stream()\n"
    "    .filter(result -> result.getStatus().equals(\"FAILED\"))\n"
    "    .map(result -> result.getTestName())\n"
    "    .collect(Collectors.toList());\n\n"
    "// WebDriverWait with lambda (Java 8 FluentWait)\n"
    "WebElement element = wait.until(driver -> {\n"
    "    WebElement el = driver.findElement(By.id(\"submit\"));\n"
    "    return el.isDisplayed() ? el : null;\n"
    "});\n\n"
    "// Optional to avoid NPE\n"
    "Optional<String> value = Optional.ofNullable(config.get(\"timeout\"));\n"
    "int timeout = Integer.parseInt(value.orElse(\"10\"));")

page_break(doc)

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 3: SELENIUM
# ═══════════════════════════════════════════════════════════════════════════════
add_section_heading(doc, "Section 3: Selenium")

add_question_heading(doc, "Q1. Explain your framework architecture")
add_para(doc, "Answer:", bold=True)
add_para(doc,
    "autoFrameX is a hybrid test automation framework built on Selenium 4 + TestNG + RestAssured. "
    "It follows a layered architecture:")
add_bullet(doc, "src/main/java/com/autoframex/core - BrowserFactory, DriverManager, BasePage, BaseTest")
add_bullet(doc, "src/main/java/com/autoframex/pages - Page Object classes (one per page/module)")
add_bullet(doc, "src/main/java/com/autoframex/api - RestAssured request/response specs, API utilities")
add_bullet(doc, "src/main/java/com/autoframex/utils - ExcelReader, JsonReader, ConfigReader, WaitUtils, ScreenshotUtils")
add_bullet(doc, "src/main/java/com/autoframex/listeners - TestNG ITestListener for ExtentReports integration")
add_bullet(doc, "src/test/java/tests - Test classes extending BaseTest")
add_bullet(doc, "src/test/resources - testng.xml suites, test data (Excel/JSON), config.properties")
add_para(doc, "Key design decisions:", bold=True)
add_bullet(doc, "ThreadLocal WebDriver for parallel execution safety")
add_bullet(doc, "Page Factory with @FindBy for element initialization")
add_bullet(doc, "Data-driven via @DataProvider reading from Excel/JSON")
add_bullet(doc, "ExtentReports with screenshots on failure")
add_bullet(doc, "Jenkins pipeline with parameterized browser and environment selection")

add_question_heading(doc, "Q2. Types of locators and XPath strategies")
add_para(doc, "Answer:", bold=True)
add_para(doc, "Locator priority (most stable to least stable):", bold=True)
add_bullet(doc, "1. ID - fastest, most reliable: By.id(\"username\")")
add_bullet(doc, "2. Name: By.name(\"email\")")
add_bullet(doc, "3. CSS Selector - fast, flexible: By.cssSelector(\"input[type='submit']\")")
add_bullet(doc, "4. XPath - most powerful, use when others fail")
add_bullet(doc, "5. LinkText / PartialLinkText: By.linkText(\"Sign In\")")
add_bullet(doc, "6. TagName, ClassName - avoid, too generic")
add_para(doc, "XPath strategies:", bold=True)
add_bullet(doc, "Absolute XPath: /html/body/div[1]/form - fragile, avoid")
add_bullet(doc, "Relative XPath: //input[@id='username'] - preferred")
add_bullet(doc, "Contains: //button[contains(text(),'Submit')]")
add_bullet(doc, "Starts-with: //input[starts-with(@id,'user')]")
add_bullet(doc, "Axes - Following sibling: //label[text()='Email']/following-sibling::input")
add_bullet(doc, "Axes - Ancestor: //input[@id='pwd']/ancestor::form")
add_bullet(doc, "Axes - Preceding sibling: //td[text()='John']/preceding-sibling::td")
add_code(doc,
    "// Dynamic XPath examples\n"
    "By.xpath(\"//td[contains(text(),'\" + dynamicValue + \"')]\");\n"
    "By.xpath(\"(//table[@id='results']//tr)[\" + rowIndex + \"]/td[2]\");")

add_question_heading(doc, "Q3. Waits - Implicit, Explicit, Fluent")
add_para(doc, "Answer:", bold=True)
add_bullet(doc, "Implicit Wait: global wait applied to all findElement calls. Driver polls DOM for specified duration before throwing NoSuchElementException. Set once, applies everywhere. Downside: can slow tests if element is genuinely absent.")
add_bullet(doc, "Explicit Wait (WebDriverWait): waits for a specific condition on a specific element. More precise. Uses ExpectedConditions. Preferred approach.")
add_bullet(doc, "Fluent Wait: customizable explicit wait. Can set polling interval, ignore specific exceptions. Most flexible.")
add_para(doc, "Best practice: use Explicit Wait. Avoid mixing Implicit + Explicit (can cause unpredictable timeouts).", bold=False)
add_code(doc,
    "// Implicit Wait\n"
    "driver.manage().timeouts().implicitlyWait(Duration.ofSeconds(10));\n\n"
    "// Explicit Wait\n"
    "WebDriverWait wait = new WebDriverWait(driver, Duration.ofSeconds(15));\n"
    "WebElement element = wait.until(ExpectedConditions.visibilityOfElementLocated(By.id(\"result\")));\n\n"
    "// Fluent Wait\n"
    "Wait<WebDriver> fluentWait = new FluentWait<>(driver)\n"
    "    .withTimeout(Duration.ofSeconds(30))\n"
    "    .pollingEvery(Duration.ofMillis(500))\n"
    "    .ignoring(NoSuchElementException.class);\n"
    "WebElement el = fluentWait.until(d -> d.findElement(By.id(\"dynamicEl\")));")

add_question_heading(doc, "Q4. Common Selenium Exceptions")
add_para(doc, "Answer:", bold=True)
add_bullet(doc, "NoSuchElementException: element not found in DOM. Fix: check locator, add explicit wait.")
add_bullet(doc, "StaleElementReferenceException: element was found but DOM was refreshed/updated, making the reference stale. Fix: re-find the element, or use a retry mechanism.")
add_bullet(doc, "ElementNotInteractableException: element exists but cannot be interacted with (hidden, disabled, covered). Fix: scroll to element, wait for it to be clickable, use JS executor.")
add_bullet(doc, "TimeoutException: explicit wait condition not met within timeout. Fix: increase timeout or fix the condition.")
add_bullet(doc, "WebDriverException: general driver error. Often caused by driver/browser version mismatch.")
add_bullet(doc, "NoSuchWindowException: window handle no longer valid. Fix: verify window handle before switching.")
add_code(doc,
    "// Handling StaleElementReferenceException with retry\n"
    "public void clickWithRetry(By locator, int maxRetries) {\n"
    "    for (int i = 0; i < maxRetries; i++) {\n"
    "        try {\n"
    "            driver.findElement(locator).click();\n"
    "            return;\n"
    "        } catch (StaleElementReferenceException e) {\n"
    "            if (i == maxRetries - 1) throw e;\n"
    "        }\n"
    "    }\n"
    "}")

add_question_heading(doc, "Q5. How to handle multiple windows")
add_para(doc, "Answer:", bold=True)
add_code(doc,
    "// Get all window handles\n"
    "String parentWindow = driver.getWindowHandle();\n"
    "Set<String> allWindows = driver.getWindowHandles();\n\n"
    "// Switch to new window\n"
    "for (String handle : allWindows) {\n"
    "    if (!handle.equals(parentWindow)) {\n"
    "        driver.switchTo().window(handle);\n"
    "        break;\n"
    "    }\n"
    "}\n"
    "// Perform actions in new window\n"
    "// ...\n"
    "// Switch back to parent\n"
    "driver.switchTo().window(parentWindow);\n\n"
    "// Selenium 4 - new window/tab\n"
    "driver.switchTo().newWindow(WindowType.TAB);\n"
    "driver.switchTo().newWindow(WindowType.WINDOW);")

add_question_heading(doc, "Q6. How to handle frames")
add_para(doc, "Answer:", bold=True)
add_code(doc,
    "// Switch to frame by index\n"
    "driver.switchTo().frame(0);\n\n"
    "// Switch to frame by name or id\n"
    "driver.switchTo().frame(\"frameName\");\n\n"
    "// Switch to frame by WebElement\n"
    "WebElement frameEl = driver.findElement(By.cssSelector(\"iframe.content-frame\"));\n"
    "driver.switchTo().frame(frameEl);\n\n"
    "// Switch to nested frame\n"
    "driver.switchTo().frame(\"outerFrame\");\n"
    "driver.switchTo().frame(\"innerFrame\");\n\n"
    "// Switch back to main document\n"
    "driver.switchTo().defaultContent();\n\n"
    "// Switch to parent frame (one level up)\n"
    "driver.switchTo().parentFrame();")

add_question_heading(doc, "Q7. Action class")
add_para(doc, "Answer:", bold=True)
add_para(doc,
    "Actions class is used to simulate complex user interactions like mouse hover, drag-and-drop, "
    "right-click, double-click, keyboard combinations, and slow typing.")
add_code(doc,
    "Actions actions = new Actions(driver);\n\n"
    "// Mouse hover\n"
    "actions.moveToElement(menuItem).perform();\n\n"
    "// Right-click (context menu)\n"
    "actions.contextClick(element).perform();\n\n"
    "// Double-click\n"
    "actions.doubleClick(element).perform();\n\n"
    "// Drag and drop\n"
    "actions.dragAndDrop(source, target).perform();\n\n"
    "// Key combination (Ctrl+A)\n"
    "actions.keyDown(Keys.CONTROL).sendKeys(\"a\").keyUp(Keys.CONTROL).perform();\n\n"
    "// Chained actions\n"
    "actions.moveToElement(menu).click().moveToElement(subMenu).click().perform();")

add_question_heading(doc, "Q8. Page Object Model")
add_para(doc, "Answer:", bold=True)
add_para(doc,
    "POM is a design pattern where each web page/component is represented as a Java class. "
    "The class contains: (1) locators as private fields, (2) public methods representing user actions. "
    "Tests interact with pages through these methods, never touching locators directly.")
add_para(doc, "Benefits:", bold=True)
add_bullet(doc, "Separation of concerns: test logic vs page interaction logic")
add_bullet(doc, "Maintainability: locator change = update one place only")
add_bullet(doc, "Reusability: same page method used across multiple tests")
add_bullet(doc, "Readability: test reads like a user story")
add_code(doc,
    "// Page class\n"
    "public class LoginPage extends BasePage {\n"
    "    @FindBy(id = \"username\") private WebElement usernameField;\n"
    "    @FindBy(id = \"password\") private WebElement passwordField;\n"
    "    @FindBy(css = \"button[type='submit']\") private WebElement loginBtn;\n\n"
    "    public LoginPage(WebDriver driver) {\n"
    "        super(driver);\n"
    "        PageFactory.initElements(driver, this);\n"
    "    }\n"
    "    public DashboardPage login(String user, String pass) {\n"
    "        usernameField.sendKeys(user);\n"
    "        passwordField.sendKeys(pass);\n"
    "        loginBtn.click();\n"
    "        return new DashboardPage(driver);\n"
    "    }\n"
    "}\n\n"
    "// Test class\n"
    "@Test\n"
    "public void testLogin() {\n"
    "    DashboardPage dashboard = new LoginPage(driver).login(\"admin\", \"pass\");\n"
    "    Assert.assertTrue(dashboard.isLoaded());\n"
    "}")

add_question_heading(doc, "Q9. Selenium Grid")
add_para(doc, "Answer:", bold=True)
add_para(doc,
    "Selenium Grid allows running tests in parallel across multiple machines and browsers. "
    "Architecture: Hub (central coordinator) + Nodes (machines with browsers). "
    "Selenium 4 introduced Grid 4 with a simplified standalone mode.")
add_bullet(doc, "Hub: receives test requests, routes to appropriate node based on desired capabilities")
add_bullet(doc, "Node: registers with hub, executes tests, reports results back")
add_bullet(doc, "In autoFrameX: RemoteWebDriver connects to Grid hub URL with browser capabilities")
add_code(doc,
    "// Connecting to Selenium Grid\n"
    "ChromeOptions options = new ChromeOptions();\n"
    "options.addArguments(\"--headless\");\n"
    "WebDriver driver = new RemoteWebDriver(\n"
    "    new URL(\"http://selenium-hub:4444/wd/hub\"),\n"
    "    options\n"
    ");\n\n"
    "// testng.xml for parallel execution\n"
    "// <suite name=\"Grid Suite\" parallel=\"tests\" thread-count=\"4\">\n"
    "//   <test name=\"Chrome Test\">\n"
    "//     <parameter name=\"browser\" value=\"chrome\"/>\n"
    "//   </test>\n"
    "// </suite>")

page_break(doc)

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 4: API TESTING
# ═══════════════════════════════════════════════════════════════════════════════
add_section_heading(doc, "Section 4: API Testing (RestAssured)")

add_question_heading(doc, "Q1. HTTP methods and status codes")
add_para(doc, "Answer:", bold=True)
add_para(doc, "HTTP Methods:", bold=True)
add_bullet(doc, "GET: retrieve resource. Safe + idempotent. No request body.")
add_bullet(doc, "POST: create resource. Not idempotent (multiple calls create multiple resources).")
add_bullet(doc, "PUT: replace entire resource. Idempotent.")
add_bullet(doc, "PATCH: partial update. Not necessarily idempotent.")
add_bullet(doc, "DELETE: remove resource. Idempotent.")
add_bullet(doc, "HEAD: like GET but returns only headers, no body.")
add_bullet(doc, "OPTIONS: returns supported HTTP methods for a resource.")
add_para(doc, "Status Codes:", bold=True)
add_bullet(doc, "2xx Success: 200 OK, 201 Created, 204 No Content")
add_bullet(doc, "3xx Redirection: 301 Moved Permanently, 302 Found, 304 Not Modified")
add_bullet(doc, "4xx Client Error: 400 Bad Request, 401 Unauthorized, 403 Forbidden, 404 Not Found, 409 Conflict, 422 Unprocessable Entity, 429 Too Many Requests")
add_bullet(doc, "5xx Server Error: 500 Internal Server Error, 502 Bad Gateway, 503 Service Unavailable, 504 Gateway Timeout")

add_question_heading(doc, "Q2. PUT vs POST vs PATCH")
add_para(doc, "Answer:", bold=True)
add_bullet(doc, "POST: creates a new resource. Server assigns the ID. URI: /users. Not idempotent - calling twice creates two users.")
add_bullet(doc, "PUT: replaces the entire resource at a known URI. URI: /users/123. Idempotent - calling twice has same result. If field is omitted, it gets nulled out.")
add_bullet(doc, "PATCH: partial update. Only sends fields to be changed. URI: /users/123. More efficient than PUT for small changes.")
add_para(doc, "Example - updating email only:", bold=True)
add_code(doc,
    "// PUT - must send entire user object\n"
    "PUT /users/123\n"
    "{\"name\": \"John\", \"email\": \"new@email.com\", \"age\": 30, \"role\": \"admin\"}\n\n"
    "// PATCH - send only changed field\n"
    "PATCH /users/123\n"
    "{\"email\": \"new@email.com\"}")

add_question_heading(doc, "Q3. Authentication - OAuth2, Bearer token, Basic auth")
add_para(doc, "Answer:", bold=True)
add_bullet(doc, "Basic Auth: username:password encoded in Base64, sent in Authorization header. Simple but insecure over HTTP. Use only with HTTPS.")
add_bullet(doc, "Bearer Token: token-based auth. Client sends 'Authorization: Bearer <token>' header. Token obtained via login endpoint. Used in most modern REST APIs.")
add_bullet(doc, "OAuth2: authorization framework. Client gets access token from Authorization Server using client_id + client_secret (or user credentials). Token has scope and expiry. Used in Salesforce, MS Dynamics 365 APIs.")
add_code(doc,
    "// Basic Auth in RestAssured\n"
    "given().auth().basic(\"username\", \"password\")\n\n"
    "// Bearer Token\n"
    "given().header(\"Authorization\", \"Bearer \" + accessToken)\n\n"
    "// OAuth2 - get token first\n"
    "String token = given()\n"
    "    .formParam(\"grant_type\", \"client_credentials\")\n"
    "    .formParam(\"client_id\", clientId)\n"
    "    .formParam(\"client_secret\", clientSecret)\n"
    "    .post(\"/oauth/token\")\n"
    "    .jsonPath().getString(\"access_token\");\n\n"
    "// Use token in subsequent requests\n"
    "given().auth().oauth2(token).get(\"/api/users\");")

add_question_heading(doc, "Q4. Serialization and Deserialization / POJO")
add_para(doc, "Answer:", bold=True)
add_bullet(doc, "Serialization: converting a Java object (POJO) to JSON/XML string for sending in request body.")
add_bullet(doc, "Deserialization: converting JSON/XML response back to a Java object (POJO).")
add_bullet(doc, "POJO (Plain Old Java Object): simple Java class with fields, getters, setters. Used to model request/response bodies.")
add_para(doc, "In autoFrameX with Jackson/RestAssured:", bold=True)
add_code(doc,
    "// POJO class\n"
    "public class User {\n"
    "    private String name;\n"
    "    private String email;\n"
    "    // getters and setters\n"
    "}\n\n"
    "// Serialization - POJO to JSON in request\n"
    "User user = new User();\n"
    "user.setName(\"John\");\n"
    "user.setEmail(\"john@example.com\");\n"
    "given().contentType(ContentType.JSON).body(user).post(\"/users\");\n\n"
    "// Deserialization - JSON response to POJO\n"
    "User createdUser = given().get(\"/users/1\")\n"
    "    .then().extract().as(User.class);\n"
    "System.out.println(createdUser.getName());")

add_question_heading(doc, "Q5. RestAssured framework explanation")
add_para(doc, "Answer:", bold=True)
add_para(doc,
    "RestAssured is a Java DSL (Domain Specific Language) for testing REST APIs. "
    "It follows a Given-When-Then BDD syntax and integrates seamlessly with TestNG/JUnit.")
add_code(doc,
    "// Basic structure\n"
    "given()           // request setup (headers, auth, body, params)\n"
    "    .baseUri(\"https://api.example.com\")\n"
    "    .header(\"Authorization\", \"Bearer \" + token)\n"
    "    .contentType(ContentType.JSON)\n"
    "    .body(requestBody)\n"
    ".when()           // HTTP method + endpoint\n"
    "    .post(\"/users\")\n"
    ".then()           // assertions\n"
    "    .statusCode(201)\n"
    "    .body(\"id\", notNullValue())\n"
    "    .body(\"name\", equalTo(\"John\"));\n\n"
    "// Extract response value\n"
    "String userId = given().get(\"/users/1\")\n"
    "    .then().statusCode(200)\n"
    "    .extract().jsonPath().getString(\"id\");")

add_question_heading(doc, "Q6. Request Specification")
add_para(doc, "Answer:", bold=True)
add_para(doc,
    "RequestSpecification allows defining common request settings once and reusing them across tests. "
    "Avoids repeating base URL, headers, and auth in every test.")
add_code(doc,
    "// Define reusable spec\n"
    "public class ApiConfig {\n"
    "    public static RequestSpecification getSpec() {\n"
    "        return new RequestSpecBuilder()\n"
    "            .setBaseUri(ConfigReader.get(\"api.baseUrl\"))\n"
    "            .addHeader(\"Authorization\", \"Bearer \" + TokenManager.getToken())\n"
    "            .setContentType(ContentType.JSON)\n"
    "            .build();\n"
    "    }\n"
    "}\n\n"
    "// Use in tests\n"
    "given().spec(ApiConfig.getSpec())\n"
    "    .body(payload)\n"
    "    .when().post(\"/users\")\n"
    "    .then().statusCode(201);")

add_question_heading(doc, "Q7. Idempotent methods")
add_para(doc, "Answer:", bold=True)
add_para(doc,
    "An HTTP method is idempotent if making the same request multiple times produces the same result "
    "as making it once. The server state is the same after one or N identical requests.")
add_bullet(doc, "Idempotent: GET, HEAD, PUT, DELETE, OPTIONS")
add_bullet(doc, "NOT idempotent: POST, PATCH (in general)")
add_para(doc, "Why it matters:", bold=True)
add_bullet(doc, "Safe to retry on network failure without side effects")
add_bullet(doc, "DELETE /users/123 called twice: first call deletes user, second call returns 404 - but server state is the same (user is gone)")
add_bullet(doc, "POST /users called twice: creates two users - not idempotent")

page_break(doc)

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 5: TestNG
# ═══════════════════════════════════════════════════════════════════════════════
add_section_heading(doc, "Section 5: TestNG")

add_question_heading(doc, "Q1. TestNG Annotations and their execution hierarchy")
add_para(doc, "Answer:", bold=True)
add_para(doc, "Execution order (top to bottom):", bold=True)
add_code(doc,
    "@BeforeSuite    -> runs once before all tests in the suite\n"
    "@BeforeTest     -> runs before each <test> tag in testng.xml\n"
    "@BeforeClass    -> runs once before first method in the class\n"
    "@BeforeMethod   -> runs before each @Test method\n"
    "@Test           -> the actual test method\n"
    "@AfterMethod    -> runs after each @Test method\n"
    "@AfterClass     -> runs once after all methods in the class\n"
    "@AfterTest      -> runs after each <test> tag in testng.xml\n"
    "@AfterSuite     -> runs once after all tests in the suite")
add_para(doc, "In autoFrameX:", bold=True)
add_bullet(doc, "@BeforeSuite: initialize ExtentReports, read global config")
add_bullet(doc, "@BeforeMethod: launch browser, navigate to base URL")
add_bullet(doc, "@AfterMethod: take screenshot on failure, quit driver")
add_bullet(doc, "@AfterSuite: flush ExtentReports, send email notification")

add_question_heading(doc, "Q2. DataProvider")
add_para(doc, "Answer:", bold=True)
add_para(doc,
    "@DataProvider supplies multiple sets of test data to a @Test method. "
    "The test runs once per data row. Returns Object[][] where each inner array is one test run.")
add_code(doc,
    "// DataProvider in same class\n"
    "@DataProvider(name = \"loginData\")\n"
    "public Object[][] getLoginData() {\n"
    "    return new Object[][] {\n"
    "        {\"admin@test.com\", \"Admin@123\", true},\n"
    "        {\"user@test.com\",  \"User@123\",  true},\n"
    "        {\"wrong@test.com\", \"wrongpass\", false}\n"
    "    };\n"
    "}\n\n"
    "@Test(dataProvider = \"loginData\")\n"
    "public void testLogin(String email, String password, boolean expectedResult) {\n"
    "    boolean result = loginPage.login(email, password);\n"
    "    Assert.assertEquals(result, expectedResult);\n"
    "}\n\n"
    "// DataProvider reading from Excel (autoFrameX pattern)\n"
    "@DataProvider(name = \"excelData\")\n"
    "public Object[][] getExcelData() {\n"
    "    return ExcelReader.getData(\"LoginData\", \"Sheet1\");\n"
    "}")

add_question_heading(doc, "Q3. Parallel execution in TestNG")
add_para(doc, "Answer:", bold=True)
add_para(doc, "TestNG supports parallel execution at multiple levels:", bold=True)
add_bullet(doc, "parallel=\"methods\": each @Test method runs in a separate thread")
add_bullet(doc, "parallel=\"tests\": each <test> tag in testng.xml runs in a separate thread")
add_bullet(doc, "parallel=\"classes\": each test class runs in a separate thread")
add_bullet(doc, "parallel=\"instances\": each instance of a test class runs in a separate thread")
add_code(doc,
    "<!-- testng.xml -->\n"
    "<suite name=\"Regression\" parallel=\"tests\" thread-count=\"4\">\n"
    "  <test name=\"Chrome-Login\">\n"
    "    <parameter name=\"browser\" value=\"chrome\"/>\n"
    "    <classes><class name=\"tests.LoginTest\"/></classes>\n"
    "  </test>\n"
    "  <test name=\"Firefox-Login\">\n"
    "    <parameter name=\"browser\" value=\"firefox\"/>\n"
    "    <classes><class name=\"tests.LoginTest\"/></classes>\n"
    "  </test>\n"
    "</suite>\n\n"
    "// Thread-safe driver via ThreadLocal (critical for parallel)\n"
    "// Each thread gets its own WebDriver instance via DriverManager.getDriver()")

add_question_heading(doc, "Q4. TestNG Listeners")
add_para(doc, "Answer:", bold=True)
add_para(doc,
    "Listeners intercept TestNG events and allow custom actions. "
    "In autoFrameX, ITestListener is used for ExtentReports integration.")
add_bullet(doc, "ITestListener: onTestStart, onTestSuccess, onTestFailure, onTestSkipped")
add_bullet(doc, "ISuiteListener: onStart, onFinish (suite level)")
add_bullet(doc, "IRetryAnalyzer: retry failed tests automatically")
add_bullet(doc, "IAnnotationTransformer: modify annotations at runtime")
add_code(doc,
    "public class TestListener implements ITestListener {\n"
    "    @Override\n"
    "    public void onTestFailure(ITestResult result) {\n"
    "        // Take screenshot\n"
    "        String screenshotPath = ScreenshotUtils.capture(result.getName());\n"
    "        // Add to ExtentReport\n"
    "        ExtentManager.getTest().fail(result.getThrowable())\n"
    "            .addScreenCaptureFromPath(screenshotPath);\n"
    "    }\n"
    "    @Override\n"
    "    public void onTestSuccess(ITestResult result) {\n"
    "        ExtentManager.getTest().pass(\"Test passed\");\n"
    "    }\n"
    "}\n\n"
    "// Register in testng.xml\n"
    "// <listeners>\n"
    "//   <listener class-name=\"listeners.TestListener\"/>\n"
    "// </listeners>")

add_question_heading(doc, "Q5. Groups and priorities in TestNG")
add_para(doc, "Answer:", bold=True)
add_para(doc, "Groups:", bold=True)
add_bullet(doc, "Tag tests with logical categories: smoke, regression, sanity, api")
add_bullet(doc, "Run specific groups via testng.xml or Maven Surefire plugin")
add_bullet(doc, "A test can belong to multiple groups")
add_para(doc, "Priority:", bold=True)
add_bullet(doc, "Controls execution order within a class. Lower number = runs first.")
add_bullet(doc, "Default priority is 0. Negative values are valid.")
add_bullet(doc, "Tests with same priority run in alphabetical order.")
add_code(doc,
    "@Test(groups = {\"smoke\", \"regression\"}, priority = 1)\n"
    "public void testLogin() { ... }\n\n"
    "@Test(groups = {\"regression\"}, priority = 2, dependsOnMethods = \"testLogin\")\n"
    "public void testDashboard() { ... }\n\n"
    "<!-- testng.xml - run only smoke group -->\n"
    "<groups>\n"
    "  <run><include name=\"smoke\"/></run>\n"
    "</groups>")

page_break(doc)

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 6: CI/CD & GIT
# ═══════════════════════════════════════════════════════════════════════════════
add_section_heading(doc, "Section 6: CI/CD & Git")

add_question_heading(doc, "Q1. Jenkins pipeline structure in your project")
add_para(doc, "Answer:", bold=True)
add_para(doc,
    "In autoFrameX, we use a declarative Jenkins pipeline (Jenkinsfile) with parameterized builds "
    "so testers can choose browser, environment, and test suite at runtime.")
add_code(doc,
    "pipeline {\n"
    "    agent any\n"
    "    parameters {\n"
    "        choice(name: 'BROWSER', choices: ['chrome','firefox','edge'], description: 'Browser')\n"
    "        choice(name: 'ENV', choices: ['qa','staging','prod'], description: 'Environment')\n"
    "        string(name: 'SUITE', defaultValue: 'regression', description: 'TestNG suite name')\n"
    "    }\n"
    "    stages {\n"
    "        stage('Checkout') {\n"
    "            steps { git branch: 'main', url: 'https://github.com/org/autoFrameX.git' }\n"
    "        }\n"
    "        stage('Build') {\n"
    "            steps { sh 'mvn clean compile -q' }\n"
    "        }\n"
    "        stage('Test') {\n"
    "            steps {\n"
    "                sh \"mvn test -Dbrowser=${params.BROWSER} -Denv=${params.ENV} -Dsuite=${params.SUITE}\"\n"
    "            }\n"
    "        }\n"
    "        stage('Reports') {\n"
    "            steps { publishHTML(target: [reportDir: 'target/extent-reports', reportFiles: 'index.html']) }\n"
    "        }\n"
    "    }\n"
    "    post {\n"
    "        failure { emailext subject: 'Build Failed', body: '${BUILD_URL}', to: 'team@company.com' }\n"
    "        always  { junit 'target/surefire-reports/*.xml' }\n"
    "    }\n"
    "}")

add_question_heading(doc, "Q2. Git commands used daily")
add_para(doc, "Answer:", bold=True)
add_code(doc,
    "# Daily workflow\n"
    "git status                          # check working tree state\n"
    "git pull origin main                # sync with remote\n"
    "git checkout -b feature/JIRA-123    # create feature branch\n"
    "git add src/test/java/tests/LoginTest.java\n"
    "git commit -m \"feat: add login automation for JIRA-123\"\n"
    "git push origin feature/JIRA-123\n\n"
    "# Useful commands\n"
    "git log --oneline --graph           # visual branch history\n"
    "git stash                           # save uncommitted changes temporarily\n"
    "git stash pop                       # restore stashed changes\n"
    "git diff HEAD~1                     # diff with previous commit\n"
    "git cherry-pick <commit-hash>       # apply specific commit to current branch\n"
    "git revert <commit-hash>            # undo a commit safely (creates new commit)\n"
    "git reset --soft HEAD~1             # undo last commit, keep changes staged")

add_question_heading(doc, "Q3. Merge conflict resolution")
add_para(doc, "Answer:", bold=True)
add_star_para(doc, "Situation:", "Two developers modified the same BasePage.java file on different branches.")
add_star_para(doc, "Task:", "Merge both branches without losing either set of changes.")
add_star_para(doc, "Action:",
    "1. git fetch origin && git merge origin/main\n"
    "2. Git marks conflict zones with <<<<<<, =======, >>>>>>>\n"
    "3. Open the file in IDE - IntelliJ shows a 3-way merge view\n"
    "4. Manually review both changes and decide: keep mine, keep theirs, or combine both\n"
    "5. Remove conflict markers, save the file\n"
    "6. git add BasePage.java && git commit -m 'resolve merge conflict in BasePage'\n"
    "7. Run tests to verify nothing broke")
add_star_para(doc, "Result:", "Clean merge with all functionality preserved. Established a team rule: always pull before starting work to minimize conflicts.")

add_question_heading(doc, "Q4. Jenkins configuration - parameterized builds")
add_para(doc, "Answer:", bold=True)
add_para(doc, "In autoFrameX, Jenkins parameters are passed to Maven as system properties:", bold=False)
add_code(doc,
    "// Reading Jenkins parameters in framework\n"
    "public class ConfigReader {\n"
    "    public static String getBrowser() {\n"
    "        // System property from Jenkins overrides config file\n"
    "        return System.getProperty(\"browser\",\n"
    "               properties.getProperty(\"browser\", \"chrome\"));\n"
    "    }\n"
    "    public static String getEnv() {\n"
    "        return System.getProperty(\"env\",\n"
    "               properties.getProperty(\"env\", \"qa\"));\n"
    "    }\n"
    "}\n\n"
    "// Maven command from Jenkins\n"
    "mvn test -Dbrowser=chrome -Denv=staging -Dsuite=regression\n\n"
    "// pom.xml surefire plugin\n"
    "<plugin>\n"
    "  <groupId>org.apache.maven.plugins</groupId>\n"
    "  <artifactId>maven-surefire-plugin</artifactId>\n"
    "  <configuration>\n"
    "    <suiteXmlFiles>\n"
    "      <suiteXmlFile>src/test/resources/${suite}.xml</suiteXmlFile>\n"
    "    </suiteXmlFiles>\n"
    "  </configuration>\n"
    "</plugin>")

page_break(doc)

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 7: DESIGN PATTERNS
# ═══════════════════════════════════════════════════════════════════════════════
add_section_heading(doc, "Section 7: Design Patterns")

add_question_heading(doc, "Q1. Singleton Pattern - with code and project usage")
add_para(doc, "Answer:", bold=True)
add_para(doc,
    "Singleton ensures only one instance of a class exists throughout the application lifecycle. "
    "In autoFrameX, it is used for ConfigReader (one config loaded once) and ExtentReports manager "
    "(one report instance shared across all tests).")
add_code(doc,
    "// Thread-safe Singleton using double-checked locking\n"
    "public class ConfigReader {\n"
    "    private static volatile ConfigReader instance;\n"
    "    private Properties properties;\n\n"
    "    private ConfigReader() {\n"
    "        properties = new Properties();\n"
    "        try (InputStream is = getClass().getClassLoader()\n"
    "                .getResourceAsStream(\"config.properties\")) {\n"
    "            properties.load(is);\n"
    "        } catch (IOException e) {\n"
    "            throw new RuntimeException(\"Cannot load config\", e);\n"
    "        }\n"
    "    }\n\n"
    "    public static ConfigReader getInstance() {\n"
    "        if (instance == null) {\n"
    "            synchronized (ConfigReader.class) {\n"
    "                if (instance == null) {\n"
    "                    instance = new ConfigReader();\n"
    "                }\n"
    "            }\n"
    "        }\n"
    "        return instance;\n"
    "    }\n\n"
    "    public String get(String key) {\n"
    "        return properties.getProperty(key);\n"
    "    }\n"
    "}\n\n"
    "// Usage\n"
    "String baseUrl = ConfigReader.getInstance().get(\"base.url\");")

add_question_heading(doc, "Q2. Builder Pattern - with project usage")
add_para(doc, "Answer:", bold=True)
add_para(doc,
    "Builder pattern constructs complex objects step by step. "
    "In autoFrameX, it is used for building API request payloads and for RequestSpecBuilder in RestAssured.")
add_code(doc,
    "// Builder for User request payload\n"
    "public class UserPayload {\n"
    "    private String name;\n"
    "    private String email;\n"
    "    private String role;\n\n"
    "    private UserPayload() {}\n\n"
    "    public static class Builder {\n"
    "        private UserPayload payload = new UserPayload();\n"
    "        public Builder name(String name)   { payload.name = name;   return this; }\n"
    "        public Builder email(String email) { payload.email = email; return this; }\n"
    "        public Builder role(String role)   { payload.role = role;   return this; }\n"
    "        public UserPayload build()         { return payload; }\n"
    "    }\n"
    "}\n\n"
    "// Usage - readable, no constructor parameter confusion\n"
    "UserPayload user = new UserPayload.Builder()\n"
    "    .name(\"John Doe\")\n"
    "    .email(\"john@example.com\")\n"
    "    .role(\"admin\")\n"
    "    .build();\n\n"
    "// RestAssured RequestSpecBuilder (built-in builder pattern)\n"
    "RequestSpecification spec = new RequestSpecBuilder()\n"
    "    .setBaseUri(baseUrl)\n"
    "    .addHeader(\"Authorization\", \"Bearer \" + token)\n"
    "    .setContentType(ContentType.JSON)\n"
    "    .build();")

add_question_heading(doc, "Q3. SOLID Principles in the framework")
add_para(doc, "Answer:", bold=True)
add_bullet(doc,
    "S - Single Responsibility: Each class has one job. LoginPage handles only login interactions. "
    "ConfigReader only reads config. ScreenshotUtils only captures screenshots.")
add_bullet(doc,
    "O - Open/Closed: Framework is open for extension, closed for modification. "
    "Adding a new browser means adding a case in BrowserFactory, not changing existing code. "
    "New page classes extend BasePage without modifying it.")
add_bullet(doc,
    "L - Liskov Substitution: Any page class can be used wherever BasePage is expected. "
    "WebDriver reference can hold ChromeDriver, FirefoxDriver, or RemoteWebDriver interchangeably.")
add_bullet(doc,
    "I - Interface Segregation: IPage interface is small and focused (isLoaded, getTitle). "
    "Not one fat interface forcing all pages to implement irrelevant methods.")
add_bullet(doc,
    "D - Dependency Inversion: High-level test classes depend on abstractions (BasePage, IPage), "
    "not on concrete implementations. BrowserFactory returns WebDriver interface, not ChromeDriver.")

page_break(doc)

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 8: SQL
# ═══════════════════════════════════════════════════════════════════════════════
add_section_heading(doc, "Section 8: SQL")

add_question_heading(doc, "Q1. JOIN queries - types and examples")
add_para(doc, "Answer:", bold=True)
add_bullet(doc, "INNER JOIN: returns rows where there is a match in BOTH tables.")
add_bullet(doc, "LEFT JOIN (LEFT OUTER JOIN): returns ALL rows from left table + matched rows from right. NULL for unmatched right rows.")
add_bullet(doc, "RIGHT JOIN: returns ALL rows from right table + matched rows from left.")
add_bullet(doc, "FULL OUTER JOIN: returns all rows from both tables. NULL where no match.")
add_bullet(doc, "CROSS JOIN: cartesian product - every row from left combined with every row from right.")
add_code(doc,
    "-- Fetch employee names with their salary (INNER JOIN)\n"
    "SELECT e.FullName, s.Salary, s.Project\n"
    "FROM EmployeeDetails e\n"
    "INNER JOIN EmployeeSalary s ON e.EmpId = s.EmpId;\n\n"
    "-- All employees even if no salary record (LEFT JOIN)\n"
    "SELECT e.FullName, COALESCE(s.Salary, 0) AS Salary\n"
    "FROM EmployeeDetails e\n"
    "LEFT JOIN EmployeeSalary s ON e.EmpId = s.EmpId;\n\n"
    "-- Employees with their manager name (self join)\n"
    "SELECT e.FullName AS Employee, m.FullName AS Manager\n"
    "FROM EmployeeDetails e\n"
    "LEFT JOIN EmployeeDetails m ON e.ManagerId = m.EmpId;")

add_question_heading(doc, "Q2. Second highest salary query")
add_para(doc, "Answer:", bold=True)
add_code(doc,
    "-- Method 1: Using subquery\n"
    "SELECT MAX(Salary) AS SecondHighest\n"
    "FROM EmployeeSalary\n"
    "WHERE Salary < (SELECT MAX(Salary) FROM EmployeeSalary);\n\n"
    "-- Method 2: Using LIMIT/OFFSET (MySQL)\n"
    "SELECT DISTINCT Salary\n"
    "FROM EmployeeSalary\n"
    "ORDER BY Salary DESC\n"
    "LIMIT 1 OFFSET 1;\n\n"
    "-- Method 3: Using DENSE_RANK (handles ties correctly)\n"
    "SELECT Salary FROM (\n"
    "    SELECT Salary, DENSE_RANK() OVER (ORDER BY Salary DESC) AS rnk\n"
    "    FROM EmployeeSalary\n"
    ") ranked\n"
    "WHERE rnk = 2;\n\n"
    "-- Nth highest salary (generic)\n"
    "SELECT Salary FROM (\n"
    "    SELECT Salary, DENSE_RANK() OVER (ORDER BY Salary DESC) AS rnk\n"
    "    FROM EmployeeSalary\n"
    ") ranked\n"
    "WHERE rnk = N; -- replace N with desired rank")

add_question_heading(doc, "Q3. DELETE vs DROP vs TRUNCATE")
add_para(doc, "Answer:", bold=True)
add_bullet(doc, "DELETE: DML command. Removes specific rows based on WHERE clause. Can be rolled back (transactional). Triggers fire. Slower - logs each row deletion. Table structure remains.")
add_bullet(doc, "TRUNCATE: DDL command. Removes ALL rows from table. Cannot be rolled back in most DBs. Faster - deallocates data pages. Resets identity/auto-increment counter. No WHERE clause. Table structure remains.")
add_bullet(doc, "DROP: DDL command. Removes the entire table (structure + data + indexes + constraints). Cannot be rolled back. Table no longer exists.")
add_code(doc,
    "-- DELETE - specific rows, can rollback\n"
    "DELETE FROM EmployeeSalary WHERE Salary < 10000;\n\n"
    "-- TRUNCATE - all rows, fast, resets auto-increment\n"
    "TRUNCATE TABLE TempTestResults;\n\n"
    "-- DROP - removes table entirely\n"
    "DROP TABLE IF EXISTS TempTestResults;")

page_break(doc)

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 9: DSA / CODING PROBLEMS
# ═══════════════════════════════════════════════════════════════════════════════
add_section_heading(doc, "Section 9: DSA / Coding Problems")

add_question_heading(doc, "Q1. Palindrome check - multiple methods")
add_para(doc, "Answer:", bold=True)
add_code(doc,
    "// Method 1: Reverse and compare\n"
    "public static boolean isPalindromeReverse(String s) {\n"
    "    String cleaned = s.toLowerCase().replaceAll(\"[^a-z0-9]\", \"\");\n"
    "    String reversed = new StringBuilder(cleaned).reverse().toString();\n"
    "    return cleaned.equals(reversed);\n"
    "}\n\n"
    "// Method 2: Two-pointer approach (O(n) time, O(1) space)\n"
    "public static boolean isPalindromeTwoPointer(String s) {\n"
    "    int left = 0, right = s.length() - 1;\n"
    "    while (left < right) {\n"
    "        if (s.charAt(left) != s.charAt(right)) return false;\n"
    "        left++;\n"
    "        right--;\n"
    "    }\n"
    "    return true;\n"
    "}\n\n"
    "// Method 3: Recursive\n"
    "public static boolean isPalindromeRecursive(String s, int left, int right) {\n"
    "    if (left >= right) return true;\n"
    "    if (s.charAt(left) != s.charAt(right)) return false;\n"
    "    return isPalindromeRecursive(s, left + 1, right - 1);\n"
    "}\n\n"
    "// Test\n"
    "System.out.println(isPalindromeReverse(\"racecar\")); // true\n"
    "System.out.println(isPalindromeReverse(\"hello\"));   // false")

add_question_heading(doc, "Q2. Reverse a string and reverse words in a sentence")
add_para(doc, "Answer:", bold=True)
add_code(doc,
    "// Reverse a string\n"
    "public static String reverseString(String s) {\n"
    "    return new StringBuilder(s).reverse().toString();\n"
    "}\n\n"
    "// Reverse words in a sentence\n"
    "// Input:  \"hello world java\"\n"
    "// Output: \"java world hello\"\n"
    "public static String reverseWords(String sentence) {\n"
    "    String[] words = sentence.trim().split(\"\\\\s+\");\n"
    "    StringBuilder sb = new StringBuilder();\n"
    "    for (int i = words.length - 1; i >= 0; i--) {\n"
    "        sb.append(words[i]);\n"
    "        if (i > 0) sb.append(\" \");\n"
    "    }\n"
    "    return sb.toString();\n"
    "}\n\n"
    "// Capitalize first letter of each word\n"
    "// Input:  \"hello! welcome to java\"\n"
    "// Output: \"Hello! Welcome To Java\"\n"
    "public static String capitalizeWords(String sentence) {\n"
    "    String[] words = sentence.split(\" \");\n"
    "    StringBuilder sb = new StringBuilder();\n"
    "    for (String word : words) {\n"
    "        if (!word.isEmpty()) {\n"
    "            sb.append(Character.toUpperCase(word.charAt(0)))\n"
    "              .append(word.substring(1).toLowerCase())\n"
    "              .append(\" \");\n"
    "        }\n"
    "    }\n"
    "    return sb.toString().trim();\n"
    "}")

add_question_heading(doc, "Q3. Count characters in a string (excluding spaces)")
add_para(doc, "Answer:", bold=True)
add_code(doc,
    "// Method 1: Loop\n"
    "public static Map<Character, Integer> countChars(String s) {\n"
    "    Map<Character, Integer> map = new LinkedHashMap<>();\n"
    "    for (char c : s.toCharArray()) {\n"
    "        if (c != ' ') {\n"
    "            map.put(c, map.getOrDefault(c, 0) + 1);\n"
    "        }\n"
    "    }\n"
    "    return map;\n"
    "}\n\n"
    "// Method 2: Java 8 Streams\n"
    "public static Map<Character, Long> countCharsStream(String s) {\n"
    "    return s.chars()\n"
    "        .filter(c -> c != ' ')\n"
    "        .mapToObj(c -> (char) c)\n"
    "        .collect(Collectors.groupingBy(c -> c, Collectors.counting()));\n"
    "}\n\n"
    "// Count total non-space characters\n"
    "public static int countNonSpaceChars(String s) {\n"
    "    return (int) s.chars().filter(c -> c != ' ').count();\n"
    "}\n\n"
    "// Test\n"
    "// Input: \"hello world\"\n"
    "// Output: {h=1, e=1, l=3, o=2, w=1, r=1, d=1}")

add_question_heading(doc, "Q4. Find duplicates in an array")
add_para(doc, "Answer:", bold=True)
add_code(doc,
    "// Method 1: Using HashSet\n"
    "public static List<Integer> findDuplicates(int[] arr) {\n"
    "    Set<Integer> seen = new HashSet<>();\n"
    "    List<Integer> duplicates = new ArrayList<>();\n"
    "    for (int num : arr) {\n"
    "        if (!seen.add(num)) { // add returns false if already present\n"
    "            duplicates.add(num);\n"
    "        }\n"
    "    }\n"
    "    return duplicates;\n"
    "}\n\n"
    "// Method 2: Using HashMap to count occurrences\n"
    "public static Map<Integer, Integer> findDuplicatesWithCount(int[] arr) {\n"
    "    Map<Integer, Integer> countMap = new HashMap<>();\n"
    "    for (int num : arr) {\n"
    "        countMap.put(num, countMap.getOrDefault(num, 0) + 1);\n"
    "    }\n"
    "    return countMap.entrySet().stream()\n"
    "        .filter(e -> e.getValue() > 1)\n"
    "        .collect(Collectors.toMap(Map.Entry::getKey, Map.Entry::getValue));\n"
    "}\n\n"
    "// Method 3: Find second largest in array\n"
    "public static int secondLargest(int[] arr) {\n"
    "    int first = Integer.MIN_VALUE, second = Integer.MIN_VALUE;\n"
    "    for (int num : arr) {\n"
    "        if (num > first) { second = first; first = num; }\n"
    "        else if (num > second && num != first) { second = num; }\n"
    "    }\n"
    "    return second;\n"
    "}\n\n"
    "// Test\n"
    "int[] arr = {1, 2, 5, 3, 4, 7};\n"
    "System.out.println(secondLargest(arr)); // 5")

add_question_heading(doc, "Q5. Count words in a sentence and date arithmetic")
add_para(doc, "Answer:", bold=True)
add_code(doc,
    "// Count words in a sentence\n"
    "public static int countWords(String sentence) {\n"
    "    if (sentence == null || sentence.trim().isEmpty()) return 0;\n"
    "    return sentence.trim().split(\"\\\\s+\").length;\n"
    "}\n\n"
    "// Display date 15 days from now\n"
    "import java.time.LocalDateTime;\n"
    "import java.time.format.DateTimeFormatter;\n\n"
    "public static void printFutureDate() {\n"
    "    LocalDateTime now = LocalDateTime.now();\n"
    "    LocalDateTime future = now.plusDays(15);\n"
    "    DateTimeFormatter formatter = DateTimeFormatter.ofPattern(\"dd-MM-yyyy HH:mm:ss\");\n"
    "    System.out.println(\"Current : \" + now.format(formatter));\n"
    "    System.out.println(\"15 days : \" + future.format(formatter));\n"
    "}\n\n"
    "// Define and traverse a Map\n"
    "Map<Integer, String> map = new HashMap<>();\n"
    "map.put(1, \"Ganga Daran\");\n"
    "map.put(2, \"Mohan Anu\");\n"
    "map.put(3, \"Arun Kumar\");\n\n"
    "// Traverse using entrySet\n"
    "for (Map.Entry<Integer, String> entry : map.entrySet()) {\n"
    "    System.out.println(entry.getKey() + \" -> \" + entry.getValue());\n"
    "}\n\n"
    "// Java 8 forEach\n"
    "map.forEach((k, v) -> System.out.println(k + \" -> \" + v));")

page_break(doc)

# ═══════════════════════════════════════════════════════════════════════════════
# SAVE DOCUMENT
# ═══════════════════════════════════════════════════════════════════════════════
output_path = r"D:\E Drive\Engineering\testleaf\workspace\autoFrameX\Interview_QA_STAR.docx"
doc.save(output_path)
print(f"Document saved: {output_path}")
